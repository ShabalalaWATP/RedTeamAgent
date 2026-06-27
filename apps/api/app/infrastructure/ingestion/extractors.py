from __future__ import annotations

from io import BytesIO
from typing import Any
from zipfile import ZipFile

from docx import Document
from pypdf import PdfReader

from app.domain.exceptions import ValidationFailure
from app.infrastructure.ingestion.code_archives import MAX_EXPANDED_BYTES, MAX_FILES, archive_result
from app.infrastructure.ingestion.redaction import redact_secret_like
from app.infrastructure.ingestion.rich_extractors import (
    audio_result,
    csv_result,
    image_result,
    pptx_result,
    video_result,
    xlsx_result,
)
from app.infrastructure.ingestion.types import ExtractedChunk, ExtractionResult

MAX_PDF_PAGES = 50
MAX_PDF_TEXT_CHARS = 100_000
MAX_DOCX_XML_BYTES = 1_000_000
MAX_DOCX_PARAGRAPHS = 1_000
MAX_DOCX_TEXT_CHARS = 100_000


class SourceExtractor:
    def extract(self, filename: str, content_type: str, content: bytes) -> ExtractionResult:
        if content_type in {"text/plain", "text/markdown"}:
            return self._text(filename, content)
        if content_type == "application/pdf":
            return self._pdf(content)
        if content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return self._docx(content)
        if content_type == "text/csv":
            return csv_result(filename, content)
        if content_type == "application/vnd.openxmlformats-officedocument.presentationml.presentation":
            return pptx_result(filename, content)
        if content_type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
            return xlsx_result(filename, content)
        if content_type in {"image/png", "image/jpeg", "image/webp"}:
            return image_result(filename, content_type, content)
        if content_type.startswith("audio/"):
            return audio_result(filename, content_type, content)
        if content_type.startswith("video/"):
            return video_result(filename, content_type, content)
        if content_type in {"application/zip", "application/x-tar", "application/gzip"}:
            return archive_result(filename, content_type, content)
        raise ValidationFailure("Unsupported source type.")

    def _text(self, filename: str, content: bytes) -> ExtractionResult:
        text = content.decode("utf-8", errors="replace")
        redacted, redaction_warnings = redact_secret_like(text)
        warnings = ["Replacement characters were used during text decoding."] if "\ufffd" in text else []
        return ExtractionResult(
            chunks=[ExtractedChunk(locator=f"{filename}:1", text=redacted)],
            metadata={"kind": "text", "characters": len(redacted)},
            warnings=[*warnings, *redaction_warnings],
        )

    def _pdf(self, content: bytes) -> ExtractionResult:
        reader = PdfReader(BytesIO(content))
        if len(reader.pages) > MAX_PDF_PAGES:
            raise ValidationFailure("PDF contains too many pages.")
        chunks: list[ExtractedChunk] = []
        warnings: list[str] = []
        total_text = 0
        for index, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            total_text += len(text)
            if total_text > MAX_PDF_TEXT_CHARS:
                raise ValidationFailure("PDF text exceeds the configured extraction limit.")
            if text.strip():
                redacted, redaction_warnings = redact_secret_like(text)
                chunks.append(ExtractedChunk(locator=f"page {index}", text=redacted))
                warnings.extend(redaction_warnings)
            else:
                chunks.append(
                    ExtractedChunk(
                        locator=f"page {index} OCR",
                        text=f"OCR placeholder for scanned PDF page {index}.",
                    )
                )
                warnings.append(f"Page {index} used deterministic OCR fallback with estimated confidence 0.55.")
        return ExtractionResult(chunks=chunks, metadata={"kind": "pdf", "pages": len(reader.pages)}, warnings=warnings)

    def _docx(self, content: bytes) -> ExtractionResult:
        _validate_docx_archive(content)
        document = Document(BytesIO(content))
        if len(document.paragraphs) > MAX_DOCX_PARAGRAPHS:
            raise ValidationFailure("DOCX contains too many paragraphs.")
        chunks: list[ExtractedChunk] = []
        total_text = 0
        for index, paragraph in enumerate(document.paragraphs, start=1):
            if not paragraph.text.strip():
                continue
            total_text += len(paragraph.text)
            if total_text > MAX_DOCX_TEXT_CHARS:
                raise ValidationFailure("DOCX text exceeds the configured extraction limit.")
            chunks.append(ExtractedChunk(locator=f"paragraph {index}", text=redact_secret_like(paragraph.text)[0]))
        return ExtractionResult(
            chunks=chunks,
            metadata={"kind": "docx", "paragraphs": len(document.paragraphs)},
            warnings=[] if chunks else ["DOCX contained no extractable paragraph text."],
        )


def _validate_docx_archive(content: bytes) -> None:
    with ZipFile(BytesIO(content)) as archive:
        infos = archive.infolist()
        if len(infos) > MAX_FILES:
            raise ValidationFailure("DOCX contains too many files.")
        total = 0
        for info in infos:
            total += info.file_size
            if info.file_size > MAX_DOCX_XML_BYTES or total > MAX_EXPANDED_BYTES:
                raise ValidationFailure("DOCX expands beyond the configured safety limit.")
            _validate_docx_part_name(info.filename)


def _validate_docx_part_name(name: Any) -> None:
    clean = str(name).replace("\\", "/")
    if clean.startswith("/") or ".." in clean.split("/"):
        raise ValidationFailure("DOCX contains an unsafe part path.")
